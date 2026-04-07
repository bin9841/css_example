from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
os.environ.setdefault("MPLCONFIGDIR", tempfile.mkdtemp(prefix="mplconfig_"))
import matplotlib.pyplot as plt
from matplotlib import patches

from src.label_builder import build_proxy_label_table_from_snapshot
from src.pipeline_runner import run_end_to_end_design
from src.scorecard_modeling import (
    build_provisional_scorecard_table,
    run_proxy_scorecard_case_analysis,
)
from src.scorecard_validation import run_proxy_validation


REPORT_DIR = Path("reports")
ASSET_DIR = REPORT_DIR / "word_assets"
SOURCE_DATASET = "retail_credit_synthetic_dataset_v3_with_exposure_flags.csv"
KEY_CUSTOMER_ID = "고객번호"
KEY_REFERENCE_MONTH = "기준년월"


def create_process_word_report(
    output_path: str | Path = REPORT_DIR / "retail_behavior_scorecard_process_wiki.docx",
    asset_dir: str | Path = ASSET_DIR,
) -> Path:
    """Create a Word report with process narrative, flowcharts, and current provisional metrics."""
    output_path = Path(output_path)
    asset_dir = Path(asset_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    asset_dir.mkdir(parents=True, exist_ok=True)

    design = run_end_to_end_design(SOURCE_DATASET)
    validation = run_proxy_validation()

    base_df = pd.read_csv(SOURCE_DATASET, encoding="utf-8-sig")
    feature_df = design["feature_table_preview"]
    population_df = design["population"]
    filtered_base = base_df.merge(
        population_df[[KEY_CUSTOMER_ID, KEY_REFERENCE_MONTH]],
        on=[KEY_CUSTOMER_ID, KEY_REFERENCE_MONTH],
        how="inner",
        validate="one_to_one",
    )
    label_df = build_proxy_label_table_from_snapshot(filtered_base)
    modeling = run_proxy_scorecard_case_analysis(feature_df=feature_df, label_df=label_df)
    points_table = build_provisional_scorecard_table(modeling["woe_maps"], modeling["coefficients"])

    pipeline_image = asset_dir / "pipeline_overview.png"
    execution_image = asset_dir / "execution_flow.png"
    timeline_image = asset_dir / "split_timeline.png"
    _draw_pipeline_overview(pipeline_image)
    _draw_execution_flow(execution_image)
    _draw_split_timeline(timeline_image)

    document = Document()
    _configure_document(document)
    _add_title(document)
    _add_overview(document, design, validation)
    _add_image_section(
        document=document,
        title="1. Pipeline Overview",
        body=(
            "The first figure shows the full scorecard redevelopment pipeline. "
            "The second figure shows the currently executable flow, including the proxy label fallback."
        ),
        images=[pipeline_image, execution_image],
        widths=[6.5, 6.5],
    )
    _add_image_section(
        document=document,
        title="2. Time Structure And Split",
        body=(
            "The project follows a time-separated behavior scorecard design. "
            "Reference months are separated into validation_1, development, and validation_2_oot."
        ),
        images=[timeline_image],
        widths=[6.5],
    )
    _add_population_section(document, design)
    _add_time_bad_section(document, design)
    _add_feature_section(document, design)
    _add_modeling_section(document, modeling)
    _add_validation_section(document, validation)
    _add_points_section(document, points_table)
    _add_module_map_section(document)
    _add_assumption_section(document)
    document.save(output_path)
    return output_path


def _configure_document(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def _add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Retail Behavior Scorecard Process Wiki")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Process explanation, executable pipeline flow, and provisional model validation summary"
    ).italic = True


def _add_overview(
    document: Document,
    design: dict[str, object],
    validation: dict[str, pd.DataFrame],
) -> None:
    document.add_heading("Executive Summary", level=1)
    population_df = design["population"]
    metric_summary = validation["metric_summary"]
    dev_row = metric_summary.loc[metric_summary["period_name"] == "development"].iloc[0]
    val1_row = metric_summary.loc[metric_summary["period_name"] == "validation1"].iloc[0]
    val2_row = metric_summary.loc[metric_summary["period_name"] == "validation2_oot"].iloc[0]

    bullets = [
        f"Current development population rows: {len(population_df):,}",
        "Reference months and split design: 2022-09 validation_1, 2023-09 development, 2024-09 validation_2_oot",
        "Modeling method: development-fitted WoE transformation and logistic regression",
        "Current labeling basis: proxy snapshot label only, not forward 12M borrower-level bad event",
        f"Development discriminatory power: KS {dev_row['ks']:.4f}, ROC {dev_row['roc']:.4f}, AR {dev_row['ar']:.4f}",
        f"Validation stability: PSI validation_1 {val1_row['psi']:.4f}, validation_2_oot {val2_row['psi']:.4f}",
    ]
    for item in bullets:
        document.add_paragraph(item, style="List Bullet")


def _add_image_section(
    document: Document,
    title: str,
    body: str,
    images: Iterable[Path],
    widths: Iterable[float],
) -> None:
    document.add_heading(title, level=1)
    document.add_paragraph(body)
    for image_path, width in zip(images, widths):
        document.add_picture(str(image_path), width=Inches(width))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_population_section(document: Document, design: dict[str, object]) -> None:
    document.add_heading("3. Population Definition", level=1)
    document.add_paragraph(
        "The development population is defined at the customer-month level with key "
        "`customer_id + reference_month`. The current implementation uses only the three "
        "approved reference months and applies a waterfall exclusion structure."
    )

    exclusion_rules = [
        "reference month outside target scope",
        "non-retail exposure",
        "inactive account",
        "no valid credit card",
        "non-scoring product",
        "bad-at-observation-point customer",
    ]
    for item in exclusion_rules:
        document.add_paragraph(item, style="List Bullet")

    document.add_paragraph("Population waterfall summary:")
    _add_dataframe_table(
        document,
        design["population_waterfall"].rename(
            columns={
                "reference_month": "reference_month",
                "step_name": "step_name",
                "before_cnt": "before_cnt",
                "after_cnt": "after_cnt",
                "excluded_cnt": "excluded_cnt",
                "exclusion_rate": "exclusion_rate",
            }
        ),
        max_rows=18,
    )


def _add_time_bad_section(document: Document, design: dict[str, object]) -> None:
    document.add_heading("4. Time Structure And Bad Definition", level=1)
    document.add_paragraph(
        "Observation windows are fixed at 1M, 3M, 6M, and 12M. "
        "The policy performance period is 12 months. The executable fallback still uses a proxy snapshot label, "
        "but the policy bad definition remains Basel-oriented."
    )

    split_summary = design["split_summary"].copy()
    _add_dataframe_table(document, split_summary, max_rows=10)

    document.add_paragraph("Policy bad definition includes:")
    for item in [
        "internal delinquency >= 60 DPD",
        "external delinquency >= 60 DPD",
        "default, charge-off, or CB default registration",
        "borrower-level recognition across all eligible facilities",
        "optional indeterminate bucket for 10 to 59 DPD",
    ]:
        document.add_paragraph(item, style="List Bullet")


def _add_feature_section(document: Document, design: dict[str, object]) -> None:
    document.add_heading("5. Raw-Data-First Feature Engineering", level=1)
    document.add_paragraph(
        "All model features are designed to be derivable from raw variables only. "
        "This preserves lineage and leakage control."
    )
    document.add_paragraph("Raw-to-feature mapping examples:")
    _add_dataframe_table(document, design["raw_to_feature_mapping"], max_rows=12)

    document.add_paragraph("Candidate feature families:")
    _add_dataframe_table(document, design["candidate_feature_list"], max_rows=20)

    document.add_paragraph("Variable screening guidance:")
    _add_dataframe_table(document, design["screening_workflow"], max_rows=10)


def _add_modeling_section(document: Document, modeling: dict[str, pd.DataFrame]) -> None:
    document.add_heading("6. Scorecard Modeling", level=1)
    document.add_paragraph(
        "The modeling sequence uses development-only coarse classing, WoE transformation, "
        "case analysis, logistic regression, and governance-oriented diagnostic review."
    )

    document.add_paragraph("Case comparison summary:")
    _add_dataframe_table(document, modeling["case_comparison"], max_rows=10)

    document.add_paragraph("Selected provisional best-case features:")
    _add_dataframe_table(document, modeling["best_features"], max_rows=15)

    document.add_paragraph("Coefficient diagnostics:")
    coef_df = modeling["coefficients"].copy()
    _add_dataframe_table(document, coef_df, max_rows=15)

    document.add_paragraph("Sign consistency review:")
    _add_dataframe_table(document, modeling["sign_check"], max_rows=15)


def _add_validation_section(document: Document, validation: dict[str, pd.DataFrame]) -> None:
    document.add_heading("7. Validation Summary", level=1)
    document.add_paragraph(
        "Validation compares development, validation_1, and validation_2_oot using scorecard metrics. "
        "The interpretation is provisional because the current label is a snapshot proxy."
    )

    document.add_paragraph("Metric summary:")
    _add_dataframe_table(document, validation["metric_summary"], max_rows=10)

    document.add_paragraph("Regulatory interpretability review:")
    _add_dataframe_table(document, validation["interpretability_review"], max_rows=15)

    document.add_paragraph("Model risk summary:")
    _add_dataframe_table(document, validation["risk_summary"], max_rows=10)


def _add_points_section(document: Document, points_table: pd.DataFrame) -> None:
    document.add_heading("8. Provisional Score Table", level=1)
    document.add_paragraph(
        "The current points table is derived from the provisional best-case logistic model with PDO 40 "
        "and anchor score 500. It is a process demonstration artifact, not a final production scorecard."
    )
    summary = (
        points_table.groupby("feature_name", dropna=False)
        .agg(
            bin_cnt=("bin", "size"),
            total_cnt=("total_cnt", "sum"),
            bad_cnt=("bad_cnt", "sum"),
            good_cnt=("good_cnt", "sum"),
            iv=("iv_component", "sum"),
            min_points=("bin_points", "min"),
            max_points=("bin_points", "max"),
        )
        .reset_index()
        .sort_values("feature_name")
    )
    _add_dataframe_table(document, summary, max_rows=20)

    document.add_paragraph("Bin-level score table sample:")
    _add_dataframe_table(document, points_table, max_rows=25)


def _add_module_map_section(document: Document) -> None:
    document.add_heading("9. Code Module Map", level=1)
    rows = [
        ["src/population_definition.py", "development population and waterfall"],
        ["src/time_structure.py", "observation, performance, and split definitions"],
        ["src/bad_definition.py", "policy bad-definition logic"],
        ["src/label_builder.py", "proxy label fallback and label schema"],
        ["src/base_dataset_schema.py", "raw schema, domains, and raw-to-feature mapping"],
        ["src/candidate_features.py", "behavior feature generation from raw variables"],
        ["src/variable_screening.py", "fine classing, IV, correlation, representative selection"],
        ["src/scorecard_modeling.py", "WoE, logistic regression, case analysis, score table"],
        ["src/scorecard_validation.py", "provisional scorecard validation metrics"],
        ["src/word_report.py", "Word report and pipeline figure generation"],
    ]
    _add_simple_table(document, ["module", "role"], rows)


def _add_assumption_section(document: Document) -> None:
    document.add_heading("10. Key Assumptions And Limitations", level=1)
    notes = [
        "The current executable label is a proxy snapshot label, not a forward 12M borrower-level bad flag.",
        "The current scorecard is therefore suitable for process demonstration, code review, and provisional analytics only.",
        "All feature engineering is intentionally raw-data-first and reference-month aligned.",
        "Development / validation / OOT separation is time-based, not random.",
        "The final regulatory package should be refreshed once real forward performance events become available.",
    ]
    for item in notes:
        document.add_paragraph(item, style="List Bullet")


def _add_dataframe_table(document: Document, df: pd.DataFrame, max_rows: int = 20) -> None:
    preview = df.head(max_rows).copy()
    preview = preview.fillna("")
    preview.columns = [str(col) for col in preview.columns]
    table = document.add_table(rows=1, cols=len(preview.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, col in enumerate(preview.columns):
        table.rows[0].cells[idx].text = col
    for _, row in preview.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row.tolist()):
            if isinstance(value, float):
                cells[idx].text = f"{value:.6f}"
            else:
                cells[idx].text = str(value)
    if len(df) > max_rows:
        document.add_paragraph(f"Displayed {max_rows} rows out of {len(df):,}.")


def _add_simple_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def _draw_pipeline_overview(output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 4))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 4)
    ax.axis("off")
    labels = [
        "Raw Data",
        "Population",
        "Time / Bad",
        "Features",
        "Screening",
        "Modeling",
        "Validation",
        "Reporting",
    ]
    x_positions = [0.3, 1.9, 3.5, 5.1, 6.7, 8.3, 9.9, 11.5]
    colors = ["#DCE6F2", "#D9EAD3", "#FCE5CD", "#FFF2CC", "#EAD1DC", "#D9D2E9", "#CFE2F3", "#D0E0E3"]
    for x, label, color in zip(x_positions, labels, colors):
        box = patches.FancyBboxPatch(
            (x, 1.4),
            1.2,
            1.0,
            boxstyle="round,pad=0.08,rounding_size=0.08",
            linewidth=1.2,
            edgecolor="#3A3A3A",
            facecolor=color,
        )
        ax.add_patch(box)
        ax.text(x + 0.6, 1.9, label, ha="center", va="center", fontsize=10, weight="bold")
    for idx in range(len(x_positions) - 1):
        ax.annotate(
            "",
            xy=(x_positions[idx + 1] - 0.08, 1.9),
            xytext=(x_positions[idx] + 1.22, 1.9),
            arrowprops=dict(arrowstyle="->", lw=1.5, color="#4F4F4F"),
        )
    ax.text(7.0, 3.2, "Retail Behavior Scorecard Redevelopment Pipeline", ha="center", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def _draw_execution_flow(output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")

    nodes = [
        (5.0, 9.0, "CSV Snapshot"),
        (5.0, 7.8, "run_population_definition"),
        (5.0, 6.6, "generate_candidate_features"),
        (5.0, 5.4, "build_proxy_label_table"),
        (5.0, 4.2, "run_proxy_scorecard_case_analysis"),
        (5.0, 3.0, "run_proxy_validation"),
        (5.0, 1.8, "Excel / Word Reports"),
    ]
    for x, y, label in nodes:
        box = patches.FancyBboxPatch(
            (x - 1.9, y - 0.35),
            3.8,
            0.7,
            boxstyle="round,pad=0.08,rounding_size=0.08",
            linewidth=1.2,
            edgecolor="#2F2F2F",
            facecolor="#F4F7FB",
        )
        ax.add_patch(box)
        ax.text(x, y, label, ha="center", va="center", fontsize=10)
    for idx in range(len(nodes) - 1):
        ax.annotate(
            "",
            xy=(5.0, nodes[idx + 1][1] + 0.4),
            xytext=(5.0, nodes[idx][1] - 0.4),
            arrowprops=dict(arrowstyle="->", lw=1.4, color="#3E3E3E"),
        )
    ax.text(5.0, 9.7, "Current Executable Flow", ha="center", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def _draw_split_timeline(output_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 2.8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 2)
    ax.axis("off")
    segments = [
        (1.0, 2.5, "#CFE2F3", "2022-09\nValidation 1"),
        (4.2, 2.5, "#D9EAD3", "2023-09\nDevelopment"),
        (7.4, 2.5, "#FCE5CD", "2024-09\nValidation 2 / OOT"),
    ]
    for x, width, color, label in segments:
        rect = patches.FancyBboxPatch(
            (x, 0.7),
            width,
            0.6,
            boxstyle="round,pad=0.06,rounding_size=0.06",
            linewidth=1.0,
            edgecolor="#4A4A4A",
            facecolor=color,
        )
        ax.add_patch(rect)
        ax.text(x + width / 2, 1.0, label, ha="center", va="center", fontsize=10, weight="bold")
    ax.annotate("", xy=(10.8, 1.0), xytext=(0.6, 1.0), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.text(5.8, 1.55, "Behavior Scorecard Time Split", ha="center", fontsize=13, weight="bold")
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


if __name__ == "__main__":
    path = create_process_word_report()
    print(path)
